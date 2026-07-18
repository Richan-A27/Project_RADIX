"""Resume text extraction helpers."""

from __future__ import annotations

from pathlib import Path
from typing import Iterable

import pdfplumber
from docx import Document


class UnsupportedFileTypeError(ValueError):
    """Raised when the provided resume file type is unsupported."""


class ResumeTextExtractionError(RuntimeError):
    """Raised when text cannot be extracted from the resume file."""


def _clean_text(lines: Iterable[str]) -> str:
    """Normalize extracted text into a single readable string."""
    cleaned_lines = [line.strip().replace("\r", "") for line in lines if line and line.strip()]
    return "\n".join(cleaned_lines).strip()


def _extract_pdf_text(filepath: Path) -> str:
    """Extract text from a PDF file page by page."""
    extracted_lines: list[str] = []

    with pdfplumber.open(filepath) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text() or ""
            if page_text.strip():
                extracted_lines.append(page_text)

    return _clean_text(extracted_lines)


def _extract_docx_text(filepath: Path) -> str:
    """Extract text from a DOCX file paragraph by paragraph."""
    document = Document(str(filepath))

    def iter_docx_text() -> Iterable[str]:
        for paragraph in document.paragraphs:
            yield paragraph.text

        for table in document.tables:
            for row in table.rows:
                yield " ".join(cell.text for cell in row.cells)

    return _clean_text(iter_docx_text())


def _extract_txt_text(filepath: Path) -> str:
    """Extract text from a plain text file."""
    content = filepath.read_text(encoding="utf-8", errors="ignore")
    return content.strip()


def extract_text(filepath: str) -> str:
    """Extract and normalize text from PDF, DOCX, or TXT resume files.

    Args:
        filepath: Path to the resume file.

    Returns:
        A single cleaned text string.

    Raises:
        FileNotFoundError: If the file does not exist.
        UnsupportedFileTypeError: If the file extension is not supported.
        ResumeTextExtractionError: If the file cannot be read or contains no text.
    """
    path = Path(filepath).expanduser().resolve()

    if not path.exists():
        raise FileNotFoundError(f"Resume file not found: {path}")

    suffix = path.suffix.lower()

    try:
        if suffix == ".pdf":
            extracted = _extract_pdf_text(path)
        elif suffix == ".docx":
            extracted = _extract_docx_text(path)
        elif suffix == ".txt":
            extracted = _extract_txt_text(path)
        else:
            raise UnsupportedFileTypeError(
                f"Unsupported resume file type: {path.suffix or '[no extension]'}"
            )
    except UnsupportedFileTypeError:
        raise
    except Exception as exc:  # pragma: no cover - defensive wrapper
        raise ResumeTextExtractionError(f"Failed to extract text from {path.name}: {exc}") from exc

    if not extracted.strip():
        raise ResumeTextExtractionError(f"No text could be extracted from {path.name}")

    return extracted
